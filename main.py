# 1) Создайте класс для банковского счета с методами пополнения, снятия и проверки баланса.

class Account:

    def __init__(self, name, balance):
        self.name = name
        self.balance = balance

    def put(self, amount):
        self.balance += amount

    def take(self, amount):
        if amount > self.balance:
            print("Error")
        else:
            self.balance -= amount

    def check(self):
        print(f"{self.name}'s balance is {self.balance}")

account1 = Account("Svetlana", 10000)
account1.check()
account1.put(100)
account1.check()
account1.take(20000)
account1.take(500)
account1.check()

# 2) Создайте класс для автомобиля с методами запуска, остановки и подачи сигнала.

class Car:

    def __init__(self, model, speed=0):
        self.model = model
        self.speed = speed

    def signal(self):
        print("Beep!")

    def start(self):
        self.speed += 10

    def stop(self):
        self.speed = 0

    def get_speed(self):
        print(self.speed)


car1 = Car("Tesla")
car1.signal()
car1.get_speed()
car1.start()
car1.get_speed()
car1.stop()
car1.get_speed()

# 3) Создайте класс для человека с атрибутами для имени, возраста и адреса.

class Person:

    def __init__(self, name, age, address):
        self.name = name
        self.age = age
        self.address = address


person1 = Person("Svetlana", 25, "St Petersburg")
print(person1.name, person1.age, person1.address)

# 4) Создайте класс для университета с атрибутами имени, адреса и студентов.

class University:

    def __init__(self, name, address, students):
        self.name = name
        self.address = address
        self.students = students


uni1 = University("SPBPU", "St Petersburg", 10000)
print(uni1.name, uni1.address, uni1.students)

# 5) Создайте класс для прямоугольника с атрибутами длины и ширины и методами вычисления площади и периметра.

class Rectangle:

    def __init__(self, length, width):
        self.length = length
        self.width = width

    def area(self):
        return self.length * self.width

    def perimeter(self):
        return 2 * (self.length + self.width)


rect1 = Rectangle(5, 10)
print(rect1.area(), rect1.perimeter())

# 6) Создайте класс для корзины покупок с методами добавления и удаления товаров и расчета общей стоимости.

class Product:

    def __init__(self, name, price):
        self.name = name
        self.price = price


class Cart:

    def __init__(self, owner):
        self.owner = owner
        self.products = list()

    def add(self, *args: Product):
        for arg in args:
            self.products.append(arg)

    def delete(self, p_name):
        for i in range(len(self.products)):
            if self.products[i].name == p_name:
                self.products.pop(i)
                break

    def total(self):
        count = 0
        for p in self.products:
            count += p.price
        return count

    def __str__(self):
        p_str = ""
        for p in self.products:
            p_str += p.name + ' '
        return f"{self.owner}, your products: {p_str}. Total: {self.total()}"


apple = Product("Apple", 50)
cheese = Product("Cheese", 100)
juice = Product("Juice", 80)
cart1 = Cart("Svetlana")
cart1.add(apple, cheese, juice)
print(cart1)
cart1.delete("Cheese")
print(cart1)

# 7) Создайте класс для игры с атрибутами для имени игрока, очков, жизней и уровня. Класс должен содержать методы для:
# a)	Запуска, завершения и перезапуска игры
# b)	Получения очков (при наборе определённого количества очков уровень должен увеличиться на 1)
# c)	Потери жизней (при достижении нуля жизней должна появиться надпись Game Over). При перезапуске игры жизни восстанавливаются.

class Player:
    def __init__(self, name, points=0, lives=3, level=0):
        self.name = name
        self.points = points
        self.lives = lives
        self.level = level

    def start(self):
        print("Start")

    def end(self):
        print("Game over")

    def restart(self):
        print("Restart")
        self.lives = 3
        self.points = 0
        self.level = 0

    def add_points(self):
        self.points += 25
        if self.points > 50:
            self.level += 1
            self.points = 0

    def life_loss(self):
        self.lives -= 1
        if self.lives == 0:
            self.end()


player1 = Player("Sveta")
print(player1.name, player1.level, player1.points, player1.lives)
player1.start()
player1.add_points()
print(player1.level, player1.points, player1.lives)
player1.add_points()
player1.add_points()
print(player1.level, player1.points, player1.lives)
player1.life_loss()
player1.life_loss()
player1.life_loss()
print(player1.level, player1.points, player1.lives)
player1.restart()
print(player1.level, player1.points, player1.lives)


# 8) Задание по библиотекам docx и python-docx (пример итогового файла в demo.docx): написать код,
# который будет создавать файл Word, где обязательно должны быть:
# a)	Заголовки двух или трёх разных уровней
# b)	Обычный текст, полужирный шрифт и курсив
# c)	Подчёркнутый текст
# d)	Нумерованный или ненумерованный список
# e)	Изображения

from docx import Document
from docx.shared import Inches


document = Document()
document.add_heading("Top-3 Animals I'd Rather Be", 0)
paragraph1 = document.add_paragraph('A little ')
run = paragraph1.add_run('insight ')
run.font.italic = True
run = paragraph1.add_run('into my ')
run = paragraph1.add_run('head.')
run.font.bold = True
document.add_heading("random paragraph out of nowhere", 1)
paragraph2 = document.add_paragraph("and one more", style='Intense Quote')
document.add_heading("The Top:", 2)
document.add_paragraph('a snow monkey', style='List Bullet')
document.add_paragraph('a dolphin', style='List Bullet')
document.add_paragraph('a duck', style='List Bullet')
document.add_picture('monkey.jpg', width=Inches(3))
document.add_picture('dolphin.jpg', width=Inches(3))
document.add_picture('duck.jpg', width=Inches(3))

document.save('My_doc.docx')



# 9) Задание по openpyxl: с помощью openpyxl написать код, который будет создавать какую-либо матрицу товаров
# (пример в stationery.xlsx) и в последнем столбце автоматически рассчитывать итоговую стоимость для каждой строки
# (кол-во товара*цену за единицу).

from openpyxl import load_workbook


workbook = load_workbook('my_table.xlsx')
worksheet = workbook.active

array = [['Название', 'Количество', 'Цена', 'Общая стоимость'], ["Яблоко", 10, 50], ["Апельсин", 8, 70],["Груша", 5, 60]]
for element in array:
    if len(element) != 4:
        total = element[1] * element[2]
        element.append(total)


for subarray in array:
    worksheet.append(subarray)
workbook.save('my_table.xlsx')